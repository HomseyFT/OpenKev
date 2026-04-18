"""
WeiWord — document editor module for OpenKev.

File format: .kev (HTML under the hood, versioned with <!-- kev:1.0 --> header)
"""

import os
import base64
from pathlib import Path

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QTabWidget, QTextEdit,
    QToolBar, QFileDialog, QMessageBox, QColorDialog, QComboBox,
    QSpinBox, QLabel, QApplication
)
from PySide6.QtGui import (
    QAction, QTextCharFormat, QFont, QTextCursor,
    QTextImageFormat, QFontDatabase, QKeySequence
)
from PySide6.QtCore import Qt, QSize
from PySide6.QtPrintSupport import QPrinter

from apps.kev_module import KevModule

KEV_VERSION_HEADER = "<!-- kev:1.0 -->\n"
UNTITLED_LABEL = "Untitled"


class DocumentTab(QWidget):
    """A single document editor pane hosted inside WeiWord's tab bar."""

    def __init__(self, filepath: str | None = None, parent=None):
        super().__init__(parent)
        self.filepath: str | None = filepath
        self._saved: bool = filepath is not None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        self.editor = QTextEdit()
        self.editor.setAcceptRichText(True)
        self.editor.textChanged.connect(self._on_text_changed)
        layout.addWidget(self.editor)

        if filepath:
            self._load(filepath)

    @property
    def is_saved(self) -> bool:
        return self._saved

    @property
    def display_name(self) -> str:
        if self.filepath:
            return Path(self.filepath).name
        return UNTITLED_LABEL

    def save(self, filepath: str | None = None) -> bool:
        target = filepath or self.filepath
        if not target:
            return False
        try:
            html = self.editor.toHtml()
            with open(target, "w", encoding="utf-8") as f:
                f.write(KEV_VERSION_HEADER)
                f.write(html)
            self.filepath = target
            self._saved = True
            return True
        except OSError:
            return False

    def export_pdf(self, filepath: str) -> bool:
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(filepath)
            self.editor.document().print_(printer)
        except Exception:
            return False
        # Qt's print_() doesn't surface failures; consider the export successful
        # iff the output file was actually produced and is non-empty.
        try:
            return os.path.exists(filepath) and os.path.getsize(filepath) > 0
        except OSError:
            return False

    def export_docx(self, filepath: str) -> bool:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:
            QMessageBox.warning(
                self,
                "Missing dependency",
                f"python-docx is required for DOCX export: {exc}",
            )
            return False
        try:
            doc = Document()
            block = self.editor.document().begin()
            while block.isValid():
                para = doc.add_paragraph()
                it = block.begin()
                while not it.atEnd():
                    frag = it.fragment()
                    if frag.isValid():
                        fmt = frag.charFormat()
                        run = para.add_run(frag.text())
                        run.bold = fmt.fontWeight() == QFont.Weight.Bold
                        run.italic = fmt.fontItalic()
                        run.underline = fmt.fontUnderline()
                        if fmt.fontSize() > 0:
                            run.font.size = Pt(fmt.fontSize())
                    it += 1
                block = block.next()
            doc.save(filepath)
            return True
        except Exception as exc:
            QMessageBox.warning(self, "Export Failed", f"DOCX export failed: {exc}")
            return False

    def insert_image(self, filepath: str) -> None:
        with open(filepath, "rb") as f:
            data = base64.b64encode(f.read()).decode("utf-8")
        ext = Path(filepath).suffix.lstrip(".").lower()
        mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
        uri = f"data:{mime};base64,{data}"
        fmt = QTextImageFormat()
        fmt.setName(uri)
        fmt.setWidth(400)
        cursor = self.editor.textCursor()
        cursor.insertImage(fmt)

    def _load(self, filepath: str) -> None:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        content = content.replace(KEV_VERSION_HEADER, "", 1)
        self.editor.setHtml(content)
        self._saved = True

    def _on_text_changed(self) -> None:
        self._saved = False


class WeiWord(KevModule):
    app_name = "Wei Word"

    def __init__(self, parent=None):
        super().__init__(parent)
        self._setup_ui()
        self.new_document()

    @property
    def open_files(self) -> list[str]:
        files = []
        for i in range(self.tab_bar.count()):
            tab = self.tab_bar.widget(i)
            if isinstance(tab, DocumentTab) and tab.filepath:
                files.append(os.path.abspath(tab.filepath))
        return files

    def focus_file(self, filepath: str) -> None:
        abs_path = os.path.abspath(filepath)
        for i in range(self.tab_bar.count()):
            tab = self.tab_bar.widget(i)
            if isinstance(tab, DocumentTab) and tab.filepath:
                if os.path.abspath(tab.filepath) == abs_path:
                    self.tab_bar.setCurrentIndex(i)
                    return

    def new_document(self) -> None:
        tab = DocumentTab(parent=self.tab_bar)
        self._add_tab(tab)

    def open_document(self, filepath: str | None = None) -> None:
        if not filepath:
            filepath, _ = QFileDialog.getOpenFileName(
                self, "Open Document", "",
                "Kev Documents (*.kev);;All Files (*)"
            )
        if not filepath:
            return
        abs_path = os.path.abspath(filepath)
        if abs_path in self.open_files:
            self.focus_file(abs_path)
            return
        tab = DocumentTab(filepath=abs_path, parent=self.tab_bar)
        self._add_tab(tab)

    def save_document(self) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        if not tab.filepath:
            self._save_as()
            return
        if not tab.save():
            QMessageBox.warning(self, "Save Failed", "Could not save the document.")
        else:
            self._refresh_tab_title(self.tab_bar.currentIndex())

    def _save_as(self) -> bool:
        tab = self._current_tab()
        if tab is None:
            return False
        filepath, _ = QFileDialog.getSaveFileName(
            self, "Save Document", "",
            "Kev Documents (*.kev);;All Files (*)"
        )
        if not filepath:
            return False
        if not filepath.endswith(".kev"):
            filepath += ".kev"
        if not tab.save(filepath):
            QMessageBox.warning(self, "Save Failed", "Could not save the document.")
            return False
        self._refresh_tab_title(self.tab_bar.currentIndex())
        return True

    def _export(self) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        filepath, selected_filter = QFileDialog.getSaveFileName(
            self, "Export Document", "",
            "PDF (*.pdf);;Word Document (*.docx)"
        )
        if not filepath:
            return
        if "pdf" in selected_filter.lower() or filepath.endswith(".pdf"):
            if not tab.export_pdf(filepath):
                QMessageBox.warning(
                    self, "Export Failed",
                    "PDF export produced no output. Check the destination path.",
                )
        elif "docx" in selected_filter.lower() or filepath.endswith(".docx"):
            # export_docx already surfaces its own error dialog on failure.
            tab.export_docx(filepath)

    def _insert_image(self) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        filepath, _ = QFileDialog.getOpenFileName(
            self, "Insert Image", "",
            "Images (*.png *.jpg *.jpeg *.gif *.bmp *.webp)"
        )
        if filepath:
            tab.insert_image(filepath)

    def _setup_ui(self) -> None:
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        layout.addWidget(self._build_toolbar())

        self.tab_bar = QTabWidget()
        self.tab_bar.setTabsClosable(True)
        self.tab_bar.tabCloseRequested.connect(self._close_tab)
        self.tab_bar.currentChanged.connect(self._on_tab_changed)
        layout.addWidget(self.tab_bar)

    def _build_toolbar(self) -> QToolBar:
        toolbar = QToolBar()
        toolbar.setIconSize(QSize(16, 16))
        toolbar.setMovable(False)

        act_new = QAction("New", self)
        act_new.setShortcut(QKeySequence.StandardKey.New)
        act_new.triggered.connect(self.new_document)
        toolbar.addAction(act_new)

        act_open = QAction("Open", self)
        act_open.setShortcut(QKeySequence.StandardKey.Open)
        act_open.triggered.connect(lambda: self.open_document())
        toolbar.addAction(act_open)

        act_save = QAction("Save", self)
        act_save.setShortcut(QKeySequence.StandardKey.Save)
        act_save.triggered.connect(self.save_document)
        toolbar.addAction(act_save)

        act_save_as = QAction("Save As", self)
        act_save_as.setShortcut(QKeySequence("Ctrl+Shift+S"))
        act_save_as.triggered.connect(self._save_as)
        toolbar.addAction(act_save_as)

        act_export = QAction("Export", self)
        act_export.triggered.connect(self._export)
        toolbar.addAction(act_export)

        toolbar.addSeparator()

        self.font_combo = QComboBox()
        self.font_combo.setFixedWidth(160)
        self.font_combo.addItems(QFontDatabase.families())
        self.font_combo.currentTextChanged.connect(self._apply_font_family)
        toolbar.addWidget(QLabel(" Font: "))
        toolbar.addWidget(self.font_combo)

        self.size_spin = QSpinBox()
        self.size_spin.setRange(6, 96)
        self.size_spin.setValue(12)
        self.size_spin.setFixedWidth(50)
        self.size_spin.valueChanged.connect(self._apply_font_size)
        toolbar.addWidget(QLabel(" Size: "))
        toolbar.addWidget(self.size_spin)

        toolbar.addSeparator()

        self.act_bold = QAction("B", self)
        self.act_bold.setShortcut(QKeySequence.StandardKey.Bold)
        self.act_bold.setCheckable(True)
        self.act_bold.triggered.connect(self._apply_bold)
        bold_font = self.act_bold.font()
        bold_font.setBold(True)
        self.act_bold.setFont(bold_font)
        toolbar.addAction(self.act_bold)

        self.act_italic = QAction("I", self)
        self.act_italic.setShortcut(QKeySequence.StandardKey.Italic)
        self.act_italic.setCheckable(True)
        self.act_italic.triggered.connect(self._apply_italic)
        italic_font = self.act_italic.font()
        italic_font.setItalic(True)
        self.act_italic.setFont(italic_font)
        toolbar.addAction(self.act_italic)

        self.act_underline = QAction("U", self)
        self.act_underline.setShortcut(QKeySequence.StandardKey.Underline)
        self.act_underline.setCheckable(True)
        self.act_underline.triggered.connect(self._apply_underline)
        toolbar.addAction(self.act_underline)

        toolbar.addSeparator()

        act_color = QAction("Font Color", self)
        act_color.triggered.connect(self._pick_font_color)
        toolbar.addAction(act_color)

        act_highlight = QAction("Highlight", self)
        act_highlight.triggered.connect(self._pick_highlight)
        toolbar.addAction(act_highlight)

        toolbar.addSeparator()

        act_image = QAction("Insert Image", self)
        act_image.triggered.connect(self._insert_image)
        toolbar.addAction(act_image)

        return toolbar

    def _merge_fmt(self, fmt: QTextCharFormat) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        cursor = tab.editor.textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.SelectionType.WordUnderCursor)
        cursor.mergeCharFormat(fmt)
        tab.editor.mergeCurrentCharFormat(fmt)

    def _apply_bold(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontWeight(QFont.Weight.Bold if checked else QFont.Weight.Normal)
        self._merge_fmt(fmt)

    def _apply_italic(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontItalic(checked)
        self._merge_fmt(fmt)

    def _apply_underline(self, checked: bool) -> None:
        fmt = QTextCharFormat()
        fmt.setFontUnderline(checked)
        self._merge_fmt(fmt)

    def _apply_font_family(self, family: str) -> None:
        fmt = QTextCharFormat()
        fmt.setFontFamilies([family])
        self._merge_fmt(fmt)

    def _apply_font_size(self, size: int) -> None:
        fmt = QTextCharFormat()
        fmt.setFontPointSize(size)
        self._merge_fmt(fmt)

    def _pick_font_color(self) -> None:
        color = QColorDialog.getColor(Qt.GlobalColor.black, self, "Font Color")
        if color.isValid():
            fmt = QTextCharFormat()
            fmt.setForeground(color)
            self._merge_fmt(fmt)

    def _pick_highlight(self) -> None:
        color = QColorDialog.getColor(Qt.GlobalColor.yellow, self, "Highlight Color")
        if color.isValid():
            fmt = QTextCharFormat()
            fmt.setBackground(color)
            self._merge_fmt(fmt)

    def _add_tab(self, tab: DocumentTab) -> None:
        index = self.tab_bar.addTab(tab, tab.display_name)
        self.tab_bar.setCurrentIndex(index)
        tab.editor.currentCharFormatChanged.connect(self._sync_toolbar)

    def _close_tab(self, index: int) -> None:
        tab = self.tab_bar.widget(index)
        if not isinstance(tab, DocumentTab):
            return
        if not tab.is_saved:
            reply = QMessageBox.question(
                self,
                "Unsaved Changes",
                f"'{tab.display_name}' has unsaved changes. Save before closing?",
                QMessageBox.StandardButton.Save |
                QMessageBox.StandardButton.Discard |
                QMessageBox.StandardButton.Cancel,
            )
            if reply == QMessageBox.StandardButton.Cancel:
                return
            if reply == QMessageBox.StandardButton.Save:
                if not tab.filepath:
                    self.tab_bar.setCurrentIndex(index)
                    if not self._save_as():
                        return
                else:
                    tab.save()
        self.tab_bar.removeTab(index)
        tab.deleteLater()

    def _refresh_tab_title(self, index: int) -> None:
        tab = self.tab_bar.widget(index)
        if isinstance(tab, DocumentTab):
            self.tab_bar.setTabText(index, tab.display_name)

    def _current_tab(self) -> DocumentTab | None:
        tab = self.tab_bar.currentWidget()
        return tab if isinstance(tab, DocumentTab) else None

    def _on_tab_changed(self, index: int) -> None:
        tab = self.tab_bar.widget(index)
        if not isinstance(tab, DocumentTab):
            return
        self._sync_toolbar()

    def _sync_toolbar(self, fmt: QTextCharFormat | None = None) -> None:
        tab = self._current_tab()
        if tab is None:
            return
        cf = fmt or tab.editor.currentCharFormat()
        self.act_bold.setChecked(cf.fontWeight() == QFont.Weight.Bold)
        self.act_italic.setChecked(cf.fontItalic())
        self.act_underline.setChecked(cf.fontUnderline())
        families = cf.fontFamilies() or []
        if families:
            family = families[0] if isinstance(families, (list, tuple)) else str(families)
            idx = self.font_combo.findText(family)
            if idx >= 0:
                self.font_combo.blockSignals(True)
                self.font_combo.setCurrentIndex(idx)
                self.font_combo.blockSignals(False)
        if cf.fontPointSize() > 0:
            self.size_spin.blockSignals(True)
            self.size_spin.setValue(int(cf.fontPointSize()))
            self.size_spin.blockSignals(False)
